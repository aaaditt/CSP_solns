"""
Builds the BITS Pilani (Dubai Campus) Practice School-I report as an editable Word
document (internship_report.docx), following the mandated PS-I format:

  - Times New Roman throughout
  - Chapter titles: 18pt, bold, centered
  - Sub-titles (x.x): 12pt, bold
  - Body: 12pt, justified, line spacing 1.15
  - Page borders on every page except the first title page
  - Roman page numbers for front matter, Arabic (restart at 1) for chapters
  - Front matter: two title pages, abstract sheet, acknowledgements, contents (auto TOC)

Run:
    python task-report/build_report.py

Requires:
    pip install python-docx

NOTE: After opening in Word, select the Table of Contents and press F9 (Update Field)
to populate page numbers. Placeholders in [square brackets] must be filled by the student.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TNR = "Times New Roman"
CODE_FONT = "Consolas"
BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x80, 0x80, 0x80)
CODE_SHADE = "F2F2F2"


# ----------------------------------------------------------------------------
# Low-level helpers
# ----------------------------------------------------------------------------
def force_font(style, name):
    """Force a style's font across ascii / hAnsi / cs so TNR fully applies."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def page_number_field(paragraph):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = " PAGE "
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(i); run._r.append(e)


def set_pgnum(section, fmt, start=None):
    sectPr = section._sectPr
    for el in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(el)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sectPr.append(pg)


def set_border(section):
    sectPr = section._sectPr
    for el in sectPr.findall(qn("w:pgBorders")):
        sectPr.remove(el)
    pgb = OxmlElement("w:pgBorders")
    pgb.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "double")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "24")
        el.set(qn("w:color"), "auto")
        pgb.append(el)
    sectPr.append(pgb)


def footer_pagenum(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_number_field(p)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin"); r.append(b)
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'; r.append(instr)
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate"); r.append(sep)
    t = OxmlElement("w:t")
    t.text = '[Right-click here and choose "Update Field" to generate the contents with page numbers.]'
    r.append(t)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end"); r.append(end)


# ----------------------------------------------------------------------------
# Content helpers
# ----------------------------------------------------------------------------
def set_update_fields_on_open(doc):
    """Tell Word to update all fields (incl. the TOC) when the document is opened."""
    settings = doc.settings.element
    for el in settings.findall(qn("w:updateFields")):
        settings.remove(el)
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.insert(0, upd)


def body(doc, text, justify=True):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def chapter(doc, text):
    doc.add_paragraph(text, style="Heading 1")


def sub(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def centered(doc, text, size, bold=True, space_after=6, color=BLACK, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = TNR; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color
    r._element.rPr.rFonts.set(qn("w:cs"), TNR)
    return p


def label_value(doc, label, value=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label + " ")
    r.font.name = TNR; r.font.size = Pt(11); r.font.bold = True
    if value:
        r2 = p.add_run(value)
        r2.font.name = TNR; r2.font.size = Pt(11)
    return p


def shade(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_code(doc, code):
    code = code.strip("\n")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    shade(p, CODE_SHADE)
    lines = code.split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = CODE_FONT; run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:cs"), CODE_FONT)
        if i < len(lines) - 1:
            run.add_break()


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True; r.font.name = TNR; r.font.size = Pt(9); r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(10)


def logo_placeholder(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[ Insert BITS Pilani, Dubai Campus logo here ]")
    r.italic = True; r.font.name = TNR; r.font.size = Pt(11); r.font.color.rgb = GREY


# ----------------------------------------------------------------------------
# Build document
# ----------------------------------------------------------------------------
doc = Document()

# Base styles -> Times New Roman
normal = doc.styles["Normal"]
normal.font.name = TNR
normal.font.size = Pt(12)
normal.font.color.rgb = BLACK
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(6)
force_font(normal, TNR)

for sname in ("List Bullet", "List Number"):
    try:
        st = doc.styles[sname]
        st.font.name = TNR; st.font.size = Pt(12); force_font(st, TNR)
    except KeyError:
        pass

# Heading 1 -> chapter title: 18pt bold centered
h1 = doc.styles["Heading 1"]
h1.font.name = TNR; h1.font.size = Pt(18); h1.font.bold = True; h1.font.color.rgb = BLACK
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_before = Pt(12); h1.paragraph_format.space_after = Pt(14)
h1.paragraph_format.line_spacing = 1.15
force_font(h1, TNR)

# Heading 2 -> sub-title: 12pt bold
h2 = doc.styles["Heading 2"]
h2.font.name = TNR; h2.font.size = Pt(12); h2.font.bold = True; h2.font.color.rgb = BLACK
h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2.paragraph_format.space_before = Pt(10); h2.paragraph_format.space_after = Pt(4)
h2.paragraph_format.line_spacing = 1.15
force_font(h2, TNR)

# Heading 3 -> 12pt bold
h3 = doc.styles["Heading 3"]
h3.font.name = TNR; h3.font.size = Pt(12); h3.font.bold = True; h3.font.color.rgb = BLACK
h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
force_font(h3, TNR)

TITLE = "INTEGRATING AI AGENTS AND MODELS IN WEBSITES MAKING CUSTOM AGENTS"

# ============================================================
# SECTION 1 — Title page (page i, no border)
# ============================================================
sec1 = doc.sections[0]
set_pgnum(sec1, "lowerRoman", start=1)
footer_pagenum(sec1)

doc.add_paragraph()
centered(doc, "A REPORT", 12)
centered(doc, "ON", 12, space_after=18)
centered(doc, TITLE, 14, space_after=24)
centered(doc, "BY", 12, space_after=18)
centered(doc, "Aadit Chandra        2024A7PS0248U        Computer Science", 12, space_after=24)
centered(doc, "AT", 12, space_after=18)
centered(doc, "SUMMER INTERNSHIP AT CSP SOLUTIONS, DUBAI", 12, space_after=24)
centered(doc, "A Practice School – I station of", 12, space_after=12)
logo_placeholder(doc)
doc.add_paragraph()
centered(doc, "BITS Pilani, Dubai Campus", 12)
centered(doc, "Dubai International Academic City, Dubai", 12)
centered(doc, "UAE", 12, space_after=18)
centered(doc, "(JUNE 2026 – JULY 2026)", 12)

# ============================================================
# SECTION 2 — Front matter (border, roman continues)
# ============================================================
doc.add_section(WD_SECTION.NEW_PAGE)
sec2 = doc.sections[1]
set_border(sec2)
set_pgnum(sec2, "lowerRoman")

# --- Page ii: second title page ---
doc.add_paragraph()
centered(doc, "A REPORT", 12)
centered(doc, "ON", 12, space_after=18)
centered(doc, TITLE, 14, space_after=24)
centered(doc, "BY", 12, space_after=18)
centered(doc, "Aadit Chandra        2024A7PS0248U        Computer Science", 12, space_after=18)
centered(doc, "Prepared in Partial Fulfillment of the", 12)
centered(doc, "Practice School – I Course", 12, space_after=24)
centered(doc, "AT", 12, space_after=18)
centered(doc, "SUMMER INTERNSHIP AT CSP SOLUTIONS, DUBAI", 12, space_after=18)
centered(doc, "A Practice School – I station of", 12, space_after=12)
logo_placeholder(doc)
doc.add_paragraph()
centered(doc, "BITS Pilani, Dubai Campus", 12)
centered(doc, "Dubai International Academic City, Dubai", 12)
centered(doc, "UAE", 12, space_after=18)
centered(doc, "(JUNE 2026 – JULY 2026)", 12)

# --- Page iii: Abstract sheet ---
doc.add_page_break()
centered(doc, "BITS Pilani, Dubai Campus", 11)
centered(doc, "Dubai International Academic City, Dubai", 11)
centered(doc, "UAE", 11, space_after=14)

t = doc.add_table(rows=2, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = t.rows[0].cells
cells[0].paragraphs[0].add_run("Station: ").bold = True
cells[0].paragraphs[0].add_run("CSP SOLUTIONS")
cells[1].paragraphs[0].add_run("Centre: ").bold = True
cells[1].paragraphs[0].add_run("DUBAI")
cells = t.rows[1].cells
cells[0].paragraphs[0].add_run("Date of Start: ").bold = True
cells[0].paragraphs[0].add_run("10 June 2026")
cells[1].paragraphs[0].add_run("Date of Submission: ").bold = True
cells[1].paragraphs[0].add_run("31 July 2026")
for row in t.rows:
    for c in row.cells:
        for para in c.paragraphs:
            for r in para.runs:
                r.font.name = TNR; r.font.size = Pt(11)
doc.add_paragraph()

label_value(doc, "Title of the Project:",
            "Integrating AI Agents and Models in Websites, Making Custom Agents")
label_value(doc, "ID No. / Name of the student:", "2024A7PS0248U / Aadit Chandra")
label_value(doc, "Discipline of Student:", "Computer Science")
label_value(doc, "Name(s) and Designation(s) of the Expert(s):", "[Expert Name, Designation — CSP Solutions]")
label_value(doc, "Name of the PS Faculty:", "[PS Faculty Name]")
label_value(doc, "Key Words:",
            "AI Agents, Large Language Models, LangChain, Google Gemini, RAG, ReAct, Streamlit")
label_value(doc, "Project Area(s):",
            "Artificial Intelligence, Web Application Development, Software Engineering")
label_value(doc, "Abstract (Max 200 words):")

abstract = (
    "This report presents the work undertaken during a Practice School-I summer internship at "
    "CSP Solutions, Dubai, on the theme of integrating Artificial Intelligence (AI) agents and "
    "large language models (LLMs) into web applications and building custom agents. The internship "
    "was structured as three progressive tasks. The first established programming and data-handling "
    "fundamentals through a Python and pandas command-line CSV-analysis tool. The second explored "
    "the core building blocks of modern LLM applications — prompt templates, conversational "
    "memory, tool-using agents based on the ReAct paradigm, and Retrieval-Augmented Generation "
    "(RAG) — implemented using the LangChain framework with Google's Gemini model. The third and "
    "principal task combined these concepts into a complete, deployed product: a multi-tool AI agent, "
    "delivered as a Streamlit web application hosted on Streamlit Cloud, capable of autonomously "
    "selecting between a calculator, live web search, and PDF question-answering. The report describes "
    "the design, implementation, and key technical challenges of each task, together with the practical "
    "software-engineering lessons learned, and concludes with directions for future enhancement toward "
    "production-ready, website-integrated custom agents."
)
body(doc, abstract)
doc.add_paragraph()
doc.add_paragraph()
sig = doc.add_table(rows=2, cols=2)
sig.rows[0].cells[0].paragraphs[0].add_run("Signature of Student").bold = True
sig.rows[0].cells[1].paragraphs[0].add_run("Signature of PS Faculty").bold = True
sig.rows[1].cells[0].paragraphs[0].add_run("Date:").bold = True
sig.rows[1].cells[1].paragraphs[0].add_run("Date:").bold = True
for row in sig.rows:
    for c in row.cells:
        for para in c.paragraphs:
            for r in para.runs:
                r.font.name = TNR; r.font.size = Pt(11)

# --- Page iv: Acknowledgements ---
doc.add_page_break()
centered(doc, "ACKNOWLEDGEMENTS", 14, space_after=14)
body(doc,
     "I would like to express my sincere gratitude to all those who supported me throughout my "
     "Practice School-I programme and made this internship a valuable learning experience.")
body(doc,
     "I thank the Director, BITS Pilani, Dubai Campus, for providing the opportunity and the "
     "environment to undertake this Practice School-I internship.")
body(doc,
     "I am deeply grateful to [Expert Name], [Designation], and the team at CSP Solutions, Dubai, "
     "for hosting me at the station and for their guidance, mentorship, and technical insight "
     "throughout the project.")
body(doc,
     "I extend my sincere thanks to the Practice School Coordinator and to my PS-I Faculty, "
     "[PS Faculty Name], for their continuous supervision, constructive feedback, and encouragement.")
body(doc,
     "I also thank the Dean, Practice School Division, for making this programme possible.")
body(doc,
     "Finally, I am thankful to my family and friends for their constant support and motivation.")
doc.add_paragraph()
p = doc.add_paragraph(); p.add_run("Aadit Chandra").bold = True
doc.add_paragraph("2024A7PS0248U")

# --- Page v: Table of Contents ---
doc.add_page_break()
centered(doc, "CONTENTS", 14, space_after=14)
for item in ["Abstract", "Acknowledgement", "List of Figures", "List of Tables"]:
    p = doc.add_paragraph()
    r = p.add_run(item); r.font.name = TNR; r.font.size = Pt(12); r.italic = True
    p.paragraph_format.space_after = Pt(4)
doc.add_paragraph()
add_toc(doc)

# ============================================================
# SECTION 3 — Body (border, Arabic restart at 1)
# ============================================================
doc.add_section(WD_SECTION.NEW_PAGE)
sec3 = doc.sections[2]
set_border(sec3)
set_pgnum(sec3, "decimal", start=1)

# ---------- Chapter 1 ----------
chapter(doc, "Chapter 1: Introduction")
body(doc,
     "This report documents the work carried out during a Practice School-I (PS-I) summer internship "
     "at CSP Solutions, Dubai. The internship focused on Artificial Intelligence (AI) — "
     "specifically, on understanding how modern AI agents and large language models (LLMs) are built "
     "and integrated into web applications, and on developing custom AI agents that can perform useful "
     "tasks autonomously.")

sub(doc, "1.1 About the Company")
body(doc,
     "CSP Solutions is a Dubai-based technology and solutions company. [Customize this paragraph with a "
     "one- to two-line description of the company's business, its main products or services, and the "
     "team or department you worked with during the internship.]")

sub(doc, "1.2 Background: AI Agents and Large Language Models")
body(doc,
     "A large language model (LLM) is a machine-learning model trained to predict the next piece of text "
     "given some input. On its own, an LLM simply takes a list of messages as input and returns a single "
     "message as output; it has no memory, cannot access private data, and cannot perform actions such as "
     "arithmetic or searching the web. Modern AI applications are built by wrapping the model with "
     "additional capabilities: memory (by re-sending the conversation), tools (functions the model can "
     "choose to call), and retrieval (supplying relevant documents at query time). An AI agent is an LLM "
     "that can reason about a task and decide which tools to use to accomplish it.")
body(doc,
     "Integrating such agents into websites allows non-technical users to access powerful AI capabilities "
     "through a simple interface, without installing anything. This project explores that idea end to end.")

sub(doc, "1.3 Objectives of the Project")
bullet(doc, "Build a strong foundation in Python and data handling.")
bullet(doc, "Understand the core building blocks of LLM applications: prompts, memory, tools, and retrieval.")
bullet(doc, "Design and implement a custom, tool-using AI agent.")
bullet(doc, "Deliver the agent as a deployed web application accessible through a browser.")

sub(doc, "1.4 Organization of the Report")
body(doc,
     "The remainder of this report is organized as follows. Chapter 2 describes the data-handling "
     "foundation built in the first task. Chapter 3 explains the building blocks of LLM applications "
     "explored in the second task. Chapter 4 presents the design, implementation, and deployment of the "
     "multi-tool AI agent developed in the third and principal task. Chapter 5 concludes the report and "
     "outlines directions for future work.")

# ---------- Chapter 2 ----------
chapter(doc, "Chapter 2: Data Handling Fundamentals (CSV Analyzer)")
sub(doc, "2.1 Objective")
body(doc,
     "The first task was to build a command-line tool, in Python and the pandas library, that can load "
     "any CSV file and help a user explore it — producing a summary, filtering and sorting the data, "
     "and exporting the result. This task contained no AI; its purpose was to establish sound data-handling "
     "and error-handling practices that would support the later AI work.")

sub(doc, "2.2 Design and Implementation")
body(doc,
     "The tool was built in two stages: a simple version that loads and summarizes a file, and a full "
     "interactive version with a menu offering five options — show a summary, filter, sort, export "
     "to CSV, and exit. The summary reports the number of rows and columns, the column names, the data "
     "types, the count of missing values per column, and per-column numeric statistics (minimum, maximum, "
     "mean, and median).")

sub(doc, "2.3 Key Features")
body(doc,
     "A feature given particular attention is type-aware filtering. If the chosen column is numeric, the "
     "entered value is converted to a number and matched exactly; otherwise a case-insensitive string match "
     "is used. Invalid input is caught and reported rather than causing a crash:")
add_code(doc, '''if pd.api.types.is_numeric_dtype(df[column]):
    try:
        value = float(value)
    except ValueError:
        print("\\n[ERROR] That column is numeric but the value you entered is not a number.\\n")
        return df
    filtered = df[df[column] == value]
else:
    filtered = df[df[column].astype(str).str.lower() == value.lower()]''')
caption(doc, "Listing 2.1: Type-aware filtering in the CSV analyzer.")
body(doc,
     "All operations run on a working copy of the data, so the originally loaded data is never modified "
     "unless the user explicitly commits a change. File loading is wrapped in error handling for missing "
     "files, empty files, and other read errors, each returning a clear message instead of a traceback.")

sub(doc, "2.4 Outcome")
body(doc,
     "The result is a small but robust utility demonstrating the pandas fundamentals — reading data, "
     "inspecting structure, handling missing values, and transforming and exporting data — together "
     "with defensive input handling. Building it in two stages also reinforced the value of iterating from "
     "a minimal working version to a polished tool.")

# ---------- Chapter 3 ----------
chapter(doc, "Chapter 3: Building Blocks of LLM Applications")
sub(doc, "3.1 Overview")
body(doc,
     "The second task was the conceptual core of the internship. Rather than one application, three small "
     "programs were built using the LangChain framework and Google's Gemini model, each isolating one "
     "building block of LLM applications: conversational memory, tool-using agents, and Retrieval-Augmented "
     "Generation (RAG). The recurring insight is that an LLM call is simply a list of messages in and one "
     "message out; memory, tools, and retrieval are just different ways of building that message list.")

sub(doc, "3.2 Conversational Memory (Chatbot)")
body(doc,
     "Because the model itself is stateless, giving it memory means re-sending the prior conversation on "
     "every turn. LangChain's RunnableWithMessageHistory does this automatically, keyed by a session id, so "
     "several independent conversations can be kept apart:")
add_code(doc, '''chatbot = RunnableWithMessageHistory(
    chain,
    get_session_history,
    input_messages_key="input",
    history_messages_key="history",
)''')
caption(doc, "Listing 3.1: Wiring conversational memory.")

sub(doc, "3.3 Tool-Using Agents (ReAct)")
body(doc,
     "An agent gives the model the ability to take actions. Three tools were defined — a calculator, a "
     "weather lookup, and a word counter — using LangChain's @tool decorator. The model reads each "
     "tool's description to decide when to use it, and can chain tools together. Importantly, the calculator "
     "never uses Python's unsafe eval() function; instead it parses the expression into an abstract syntax "
     "tree and evaluates only a whitelist of safe operations:")
add_code(doc, '''@tool
def calculator(expression: str) -> str:
    """
    Evaluates a math expression and returns the result.
    Use this for any arithmetic: addition, subtraction, multiplication,
    division, percentages, etc. Example input: '25 * 0.25' or '100 + 50'.
    """
    try:
        tree = ast.parse(expression, mode="eval")
        result = _safe_eval(tree.body)
        return str(result)
    except Exception as e:
        return f"Error evaluating expression: {e}"''')
caption(doc, "Listing 3.2: A tool defined with the @tool decorator and safe AST-based evaluation.")

sub(doc, "3.4 Retrieval-Augmented Generation (RAG)")
body(doc,
     "RAG makes the model answer from a specific set of documents rather than its general training. At "
     "startup, the documents are split into chunks, embedded into vectors, and stored in a FAISS index. For "
     "each question, the most relevant chunks are retrieved and inserted into a prompt that instructs the "
     "model to answer using only that context — which prevents it from fabricating answers when the "
     "information is not present:")
add_code(doc, '''def ask(question: str) -> str:
    """Retrieves relevant chunks and generates an answer using the LLM."""
    relevant_chunks = retriever.invoke(question)
    ...
    context = "\\n\\n".join(chunk.page_content for chunk in relevant_chunks)
    filled_prompt = rag_prompt.invoke({"context": context, "question": question})
    response = model.invoke(filled_prompt)
    return response.content''')
caption(doc, "Listing 3.3: The retrieve-then-generate function.")

sub(doc, "3.5 Outcome")
body(doc,
     "Building the three programs separately made the concepts concrete. Memory, tools, and retrieval seem "
     "like very different features, but they are the same underlying pattern — a list of messages to "
     "the model, one message back — differing only in how the message list is assembled. This mental "
     "model directly shaped the design of the final task.")

# ---------- Chapter 4 ----------
chapter(doc, "Chapter 4: A Deployed Multi-Tool AI Agent")
sub(doc, "4.1 Objective")
body(doc,
     "The third and principal task combined the ideas from the previous task into a complete, deployed "
     "product: a web application in which an AI agent answers questions by choosing the right tool. It has "
     "four tools — a calculator, a live web search, and the ability to load a PDF and answer questions "
     "about it — and can use several of them in a single response. This directly realizes the project "
     "theme of a custom, website-integrated AI agent.")

sub(doc, "4.2 System Architecture")
body(doc,
     "The user chats through a Streamlit web interface. Each message is passed to a LangGraph ReAct agent "
     "built from the Gemini model, the four tools, and a routing prompt that tells the model what each tool "
     "is for:")
add_code(doc, '''def create_my_agent():
    """Builds the agent with Gemini + all four tools."""
    model = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        google_api_key=os.getenv("GOOGLE_API_KEY"),
        temperature=0,
    )
    tools = [calculator, web_search, load_pdf, query_pdf]
    agent = create_react_agent(
        model=model,
        tools=tools,
        prompt=(
            "You are a helpful AI assistant with access to these tools:\\n"
            "- calculator: for math expressions (percentages, arithmetic)\\n"
            "- web_search: for current information from the internet\\n"
            "- load_pdf: to load a PDF file (call this first before querying)\\n"
            "- query_pdf: to answer questions about a loaded PDF\\n\\n"
            "Pick the right tool based on the user's question. ..."
        ),
    )
    return agent''')
caption(doc, "Listing 4.1: Building the ReAct agent with Gemini and four tools.")

sub(doc, "4.3 The Web Interface")
body(doc,
     "The Streamlit chat loop captures user input, shows a spinner while the agent runs, and renders the "
     "reply, persisting the conversation in the browser session state:")
add_code(doc, '''user_input = st.chat_input("Ask me anything...")
if user_input:
    st.session_state.messages.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.markdown(user_input)
    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):
            response = run_agent(
                st.session_state.agent,
                user_input,
                st.session_state.messages[:-1],
            )
        st.markdown(response)
    st.session_state.messages.append({"role": "assistant", "content": response})
    st.rerun()''')
caption(doc, "Listing 4.2: The Streamlit chat loop.")

sub(doc, "4.4 Robustness of the Web-Search Tool")
body(doc,
     "The web-search tool wraps the Tavily search API. In practice this API can return results in several "
     "different shapes depending on the version, so the tool normalizes each of them into clean, readable "
     "text — a piece of defensive code that required real iteration to get right:")
add_code(doc, '''from langchain_tavily import TavilySearch
search = TavilySearch(max_results=3, tavily_api_key=api_key)
raw = search.invoke(query)

# TavilySearch returns different shapes depending on version...
if isinstance(raw, list):
    parts = []
    for item in raw:
        if isinstance(item, dict):
            if "text" in item:
                parts.append(item["text"])
            elif "content" in item:
                title = item.get("title", "")
                content = item.get("content", "")
                url = item.get("url", "")
                entry = f"{title}\\n{content}"
                if url:
                    entry += f"\\nSource: {url}"
                parts.append(entry)
    return "\\n\\n".join(parts) if parts else str(raw)''')
caption(doc, "Listing 4.3: Normalizing web-search results into readable text.")

sub(doc, "4.5 Deployment on Streamlit Cloud")
body(doc,
     "The application runs both locally and on Streamlit Cloud. Locally, the API keys are read from a .env "
     "file; on the cloud there is no such file, so the keys live in Streamlit's secure secrets store. A short "
     "bridge copies those secrets into environment variables at startup, so the same os.getenv() calls "
     "throughout the code work identically in both environments:")
add_code(doc, '''# Bridge Streamlit Cloud secrets -> env vars (so os.getenv works on both local and cloud)
for key in ["GOOGLE_API_KEY", "TAVILY_API_KEY"]:
    if key not in os.environ:
        try:
            os.environ[key] = st.secrets[key]
        except (KeyError, FileNotFoundError):
            pass''')
caption(doc, "Listing 4.4: The secrets bridge that lets one codebase run locally and on the cloud.")

sub(doc, "4.6 Challenges and Solutions")
bullet(doc,
       "Security — avoiding eval(): a naive calculator would pass user text to Python's eval(), a "
       "serious security risk. This was solved by parsing expressions into an abstract syntax tree and "
       "evaluating only a whitelist of arithmetic operations.")
bullet(doc,
       "Handling a messy external API: the Tavily response arrived in several different formats, so the tool "
       "was written to defensively normalize all of them into clean text.")
bullet(doc,
       "Local versus cloud configuration: deploying to Streamlit Cloud required bridging secrets into "
       "environment variables so the same code path works in both places.")
bullet(doc,
       "Preventing hallucination: in the RAG work, explicitly instructing the model to answer using only "
       "the retrieved context — and to say when it does not know — was what made its answers "
       "trustworthy.")

# ---------- Chapter 5 ----------
chapter(doc, "Chapter 5: Conclusion and Future Scope")
sub(doc, "5.1 Conclusions")
body(doc,
     "Across the three tasks, the internship progressed from core Python data handling, through the "
     "fundamental building blocks of LLM applications, to the design and deployment of a complete, "
     "website-integrated AI product. The final deliverable — a multi-tool AI agent running as a "
     "Streamlit web application — demonstrates the project theme of integrating AI agents and models "
     "into websites and building custom agents. Beyond the technical outcomes, the internship developed "
     "practical engineering habits: robust error handling, security-conscious coding, and managing the "
     "difference between local and cloud environments.")

sub(doc, "5.2 Future Scope")
bullet(doc, "Persisting conversation memory and loaded documents beyond a single process.")
bullet(doc, "Isolating state per user so the agent can serve many users safely on a shared deployment.")
bullet(doc, "Adding further tools (for example, database queries, e-mail, or calendar access).")
bullet(doc, "Embedding the agent as a widget within an existing website rather than a standalone app.")
bullet(doc, "Adding authentication, usage limits, and monitoring for a production-ready deployment.")

# ---------- References ----------
chapter(doc, "References")
refs = [
    "LangChain Documentation. https://python.langchain.com/",
    "LangGraph Documentation. https://langchain-ai.github.io/langgraph/",
    "Google Gemini API. https://ai.google.dev/",
    "Streamlit Documentation. https://docs.streamlit.io/",
    "Streamlit Community Cloud. https://docs.streamlit.io/deploy/streamlit-community-cloud",
    "Tavily Search API. https://tavily.com/",
    "pandas Documentation. https://pandas.pydata.org/docs/",
    "FAISS: A library for efficient similarity search. https://faiss.ai/",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(f"[{i}]  {r}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ---------- Turnitin ----------
doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run("Paste the Turnitin (plagiarism) report here.")
r.bold = True; r.font.name = TNR; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

set_update_fields_on_open(doc)

import sys
out_path = sys.argv[1] if len(sys.argv) > 1 else "internship_report.docx"
doc.save(out_path)
print(f"Report written to {out_path}")
